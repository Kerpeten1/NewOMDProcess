from docx import Document
import collections


# path
def fill_in_requester_data(re):
    requester_data = [re[0][1], re[1][1], re[2][1], re[3][1], re[4][1]]
    if "@" in re[5][1]:
        requester_data.append(re[5][1])

    document = Document(r"C:\Users\M261651\Desktop\Dokumente\Files für Pycharm Projekte\OMD Prozess\OMDLetter\OMD Vorlage\TCUST_standard_disclosure_letter.docx")
    document.tables[0].rows[0].cells[0].text = ""
    for element in requester_data:
        document.tables[0].rows[0].cells[0].text = document.tables[0].rows[0].cells[0].text + element + "\n"
    document.save(r"C:\Users\M261651\Desktop\Dokumente\Files für Pycharm Projekte\OMD Prozess\OMDLetter\OMD Vorlage\TCUST_standard_disclosure_letter_company_filled.docx")


# path
def fill_in_item_data(items):
    document = Document(r"C:\Users\M261651\Desktop\Dokumente\Files für Pycharm Projekte\OMD Prozess\OMDLetter\OMD Vorlage\TCUST_standard_disclosure_letter_company_filled.docx")
    row = 1
    index = 0
    for element in items:
        print("element in items: ", element)
        index += 1
        if element:
            length = len(element)
            document.tables[1].rows[row].cells[0].text = document.tables[1].rows[row].cells[0].text + element[0][0]
            document.tables[1].rows[row].cells[1].text = document.tables[1].rows[row].cells[1].text + element[0][1]
            for count_oms in range(0, length):
                document.tables[1].rows[row].cells[2].text = document.tables[1].rows[row].cells[2].text + element[count_oms][2]
                if count_oms is not length-1:
                    document.tables[1].rows[row].cells[2].add_paragraph()
                    document.tables[1].rows[row].cells[2].add_paragraph()
            add_table_row(element, items, document)
        row += 1
    document.save(r"C:\Users\M261651\Desktop\Dokumente\Files für Pycharm Projekte\OMD Prozess\OMDLetter\OMD Vorlage\TCUST_standard_disclosure_letter_company_and_items_filled.docx")


def add_table_row(element, items, document):
    element = element
    items = items
    document = document
    compare = lambda x, y: collections.Counter(x) == collections.Counter(y)
    if not(compare(element, items[-1])):
        document.tables[1].add_row()



