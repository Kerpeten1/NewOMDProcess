import docx2txt
import time
from docx import Document
from src.Fill_in_OMD import fill_in_requester_data, fill_in_item_data
from src.Find_Item_and_Manufacturer import get_rows_and_columns_of_omd_excel2
from src.Save_OMD import save_omd_as_docx


def read_rfi_data(file_name):
    # get data from rfi
    requester_data = get_requester_email(file_name)
    item_data = get_items(file_name)

    # write rfi data in omd
    fill_in_requester_data(requester_data)
    fill_in_item_data(item_data)

    # save omd
    save_omd_as_docx(requester_data[0][1])



# if email in rfi is hyperlink, then take index[74]
def get_requester_email(name):
    path = "RFIs\\"
    file_name = name
    path_and_file_name = path + file_name
    my_text = docx2txt.process(path_and_file_name)
    lines = my_text.splitlines()
    email_description = lines[71]
    email_address = lines[74]
    requester_data = get_requester(path_and_file_name, email_description, email_address)
    return requester_data


def get_requester(path_and_file_name, email, email_address):
    document = Document(path_and_file_name)
    company = document.tables[1]
    requester = []

    for row in company.rows:
        if row.cells[1].text != " " and row.cells[1].text != "":
            requester.append((row.cells[0].text, row.cells[1].text))
    if "@" not in (requester[-1][-1]):
        requester.append((email, email_address))
    return requester


def get_items(name):
    path = "RFIs\\"
    file_name = name
    path_and_file_name = path + file_name
    document = Document(path_and_file_name)
    item = document.tables[4]
    items = []

    for row in item.rows:
        if row.cells[0].text != "":
            item_number = row.cells[0].text
            item_description = row.cells[1].text
            get_manufacturer_of_item = get_manufacturer(item_number, item_description)
            # item = (item_number + "," + item_description).split(",")
            items.append(get_manufacturer_of_item)
    items = delete_none_from_item_list(items)
    # print("item, description, om: ", items)
    return items


def delete_none_from_item_list(items):
    clean = [x for x in items if x != None]
    print("Nones Deleted: ", clean)
    return clean


def get_manufacturer(item, item_description):
    manufacturer = get_rows_and_columns_of_omd_excel2(item, item_description)

    return manufacturer

